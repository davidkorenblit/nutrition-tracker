import os
import uuid
import re
import json
from fastapi import UploadFile, HTTPException
from docx import Document
import google.generativeai as genai
from supabase import create_client, Client
import logging

logger = logging.getLogger(__name__)

# Supabase Configuration
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

# Initialize Supabase client if credentials are available
supabase_client: Client = None
if SUPABASE_URL and SUPABASE_KEY:
    try:
        supabase_client = create_client(SUPABASE_URL, SUPABASE_KEY)
        logger.info("âœ… Supabase client initialized for file storage")
    except Exception as e:
        logger.warning(f"âš ï¸  Failed to initialize Supabase: {e}")

BUCKET_NAME = os.getenv("RECOMMENDATIONS_BUCKET_NAME", "recommendations")
ALLOWED_EXTENSIONS = {"docx"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB


def validate_word_file(file: UploadFile) -> bool:
    """×‘×“×™×§×ª ×ª×§×™× ×•×ª ×§×•×‘×¥ Word"""
    ext = file.filename.split(".")[-1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Use: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    return True


async def save_word_file(file: UploadFile) -> str:
    """
    ×©×ž×™×¨×ª ×§×•×‘×¥ Word ×‘-Supabase ×•×”×—×–×¨×ª public URL.
    
    Args:
        file: ×§×•×‘×¥ Word ×œ×”×¢×œ××”
    
    Returns:
        str: Public URL ×©×œ ×”×§×•×‘×¥ ×‘-Supabase
    
    Raises:
        HTTPException: ×× ×”×§×•×‘×¥ ×œ× ×ª×§×™×Ÿ ××• ×’×“×•×œ ×ž×“×™
    """
    validate_word_file(file)
    
    # ×‘×“×•×§ ×× Supabase ×ž×•×’×“×¨
    if not supabase_client:
        raise HTTPException(
            status_code=500,
            detail="Cloud storage (Supabase) is not configured. Please set SUPABASE_URL and SUPABASE_KEY."
        )
    
    # ×§×¨× ××ª ×ª×•×›×Ÿ ×”×§×•×‘×¥
    content = await file.read()
    
    # ×‘×“×•×§ ×’×•×“×œ
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Max size: {MAX_FILE_SIZE / 1024 / 1024}MB"
        )
    
    # ×¦×•×¨ ×©× ×™×™×—×•×“×™
    file_ext = file.filename.split(".")[-1]
    unique_filename = f"{uuid.uuid4()}.{file_ext}"
    
    try:
        # Upload to Supabase
        response = supabase_client.storage.from_(BUCKET_NAME).upload(
            path=unique_filename,
            file=content,
            file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
        )
        
        # Build public URL
        public_url = f"{SUPABASE_URL}/storage/v1/object/public/{BUCKET_NAME}/{unique_filename}"
        logger.info(f"âœ… File uploaded to Supabase: {unique_filename}")
        
        return public_url
    
    except Exception as e:
        logger.error(f"âŒ Failed to upload to Supabase: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to upload file to cloud storage: {str(e)}"
        )


def parse_word_file(file_path: str) -> str:
    """
    ×—×™×œ×•×¥ ×˜×§×¡×˜ ×ž×§×•×‘×¥ Word ×ž-URL ×©×œ Supabase.
    ×§×•×¨× ×’× ×¤×¡×§××•×ª ×¨×’×™×œ×•×ª ×•×’× ×ª×•×›×Ÿ ×©×œ ×˜×‘×œ××•×ª.
    
    Args:
        file_path: Public URL ×©×œ ×§×•×‘×¥ Word ×‘-Supabase
    
    Returns:
        str: ×”×˜×§×¡×˜ ×”×ž×œ× ×ž×”×§×•×‘×¥
    
    Raises:
        HTTPException: ×× ×”×§×¨×™××” × ×›×©×œ×”
    """
    try:
        import requests
        from io import BytesIO
        
        # Download file from Supabase URL
        response = requests.get(file_path)
        response.raise_for_status()
        
        # Parse using python-docx
        doc = Document(BytesIO(response.content))
        full_text = []
        
        # ×—×œ×¥ ×˜×§×¡×˜ ×ž×¤×¡×§××•×ª ×¨×’×™×œ×•×ª
        for para in doc.paragraphs:
            if para.text.strip():  # ×¨×§ ×× ×™×© ×˜×§×¡×˜
                full_text.append(para.text)
        
        # ×—×œ×¥ ×˜×§×¡×˜ ×ž×˜×‘×œ××•×ª
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # ×¨×§ ×× ×™×© ×˜×§×¡×˜
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
        
    except Exception as e:
        logger.error(f"âŒ Failed to parse Word file: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse Word file: {str(e)}"
        )


def extract_recommendations_with_llm(raw_text: str) -> list[dict]:
    """
    ×—×™×œ×•×¥ ×”×ž×œ×¦×•×ª ×ž×”×˜×§×¡×˜ ×‘××ž×¦×¢×•×ª Gemini LLM.
    
    Args:
        raw_text: ×”×˜×§×¡×˜ ×”×ž×œ× ×ž×§×•×‘×¥ Word
    
    Returns:
        list[dict]: ×¨×©×™×ž×ª ×”×ž×œ×¦×•×ª ×‘×¤×•×¨×ž×˜ [{"id": 1, "text": "..."}, ...]
    
    Raises:
        HTTPException: ×× ×”×§×¨×™××” ×œ-LLM × ×›×©×œ×”
    """
    try:
        # ×”×’×“×¨×ª API key
        api_key = os.getenv("GEMINI_API_KEY")
        if not api_key:
            raise HTTPException(
                status_code=500,
                detail="GEMINI_API_KEY not configured"
            )
        
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel('gemini-flash-latest')
        
        # ×‘× ×™×™×ª prompt
        prompt = f"""
××ª×” ×¢×•×–×¨ ×©×ž× ×ª×— ×ž×¡×ž×›×™ Word ×©×œ ×ª×–×•× ××™×.
×—×œ×¥ ××ª ×›×œ ×”×”×ž×œ×¦×•×ª ×ž×¡×¢×™×£ "×”×ž×œ×¦×•×ª ×œ×‘×™×ª" ×‘×œ×‘×“.
×”×—×–×¨ ×¨×©×™×ž×ª JSON ×‘×¤×•×¨×ž×˜ ×”×‘× ×‘×“×™×•×§, ×œ×œ× ×˜×§×¡×˜ × ×•×¡×£:
[{{"id": 1, "text": "×˜×§×¡×˜ ×”×”×ž×œ×¦×” ×”×¨××©×•× ×”"}}, {{"id": 2, "text": "×˜×§×¡×˜ ×”×”×ž×œ×¦×” ×”×©× ×™×™×”"}}]

×—×©×•×‘:
- ×¨×§ ×”×ž×œ×¦×•×ª ×ž×”×¡×¢×™×£ "×”×ž×œ×¦×•×ª ×œ×‘×™×ª"
- ×›×œ ×”×ž×œ×¦×” ×›×©×•×¨×” × ×¤×¨×“×ª
- ××œ ×ª×•×¡×™×£ ×”×¡×‘×¨×™× ××• ×˜×§×¡×˜ × ×•×¡×£, ×¨×§ JSON
- ×× ××™×Ÿ ×”×ž×œ×¦×•×ª, ×”×—×–×¨ []

×˜×§×¡×˜ ×”×ž×¡×ž×š:
{raw_text}
"""
        
        # ×§×¨×™××” ×œ-LLM
        response = model.generate_content(prompt)
        response_text = response.text.strip()
        
        # × ×™×§×•×™ ×”×ª×’×•×‘×” (×”×¡×¨×ª markdown ×× ×™×©)
        response_text = re.sub(r'```json\s*', '', response_text)
        response_text = re.sub(r'```\s*', '', response_text)
        
        # ×”×ž×¨×” ×œ-JSON
        recommendations_list = json.loads(response_text)
        
        # ×”×•×¡×£ ×©×“×•×ª ×‘×¨×™×¨×ª ×ž×—×“×œ ×œ×›×œ ×”×ž×œ×¦×”
        for rec in recommendations_list:
            rec["category"] = "general"
            rec["tracked"] = False
            rec["target_value"] = None
            rec["notes"] = None
        
        return recommendations_list
        
    except json.JSONDecodeError as e:
        logger.error(f"âŒ Failed to parse LLM response as JSON: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse LLM response as JSON: {str(e)}"
        )
    except Exception as e:
        logger.error(f"âŒ Failed to extract recommendations with LLM: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to extract recommendations with LLM: {str(e)}"
        )


def delete_word_file(file_url: str) -> bool:
    """
    ×ž×—×™×§×ª ×§×•×‘×¥ Word ×ž-Supabase.
    
    Args:
        file_url: Public URL ×©×œ ×”×§×•×‘×¥
    
    Returns:
        bool: True ×× × ×ž×—×§ ×‘×”×¦×œ×—×”, False ××—×¨×ª
    """
    if not supabase_client:
        logger.warning("âš ï¸  Supabase not configured, cannot delete file")
        return False
    
    try:
        # Extract filename from URL
        filename = file_url.split('/')[-1]
        
        # Delete from Supabase
        supabase_client.storage.from_(BUCKET_NAME).remove([filename])
        logger.info(f"âœ… File deleted from Supabase: {filename}")
        return True
    except Exception as e:
        logger.error(f"âŒ Failed to delete file from Supabase: {e}")
        return False


