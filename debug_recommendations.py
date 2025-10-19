#!/usr/bin/env python3
"""
Comprehensive test script for Recommendations and Compliance system
Tests both backend and frontend components
"""

import requests
import json
import os
import sys
from datetime import datetime, timedelta
from colorama import init, Fore, Back, Style
import time

# Initialize colorama for colored output
init(autoreset=True)

# Configuration
BASE_URL = "http://localhost:8000"
API_URL = f"{BASE_URL}/api/v1"
FRONTEND_URL = "http://localhost:3000"

# Test user credentials (dummy)
TEST_USER = {
    "email": "el@gmail.com",
    "password": "444444"
}

class TestRunner:
    def __init__(self):
        self.token = None
        self.errors = []
        self.warnings = []
        self.successes = []
        
    def print_header(self, text):
        print(f"\n{Fore.CYAN}{'='*60}")
        print(f"{Fore.CYAN}{text:^60}")
        print(f"{Fore.CYAN}{'='*60}{Style.RESET_ALL}\n")
        
    def print_success(self, message):
        print(f"{Fore.GREEN}✓ {message}{Style.RESET_ALL}")
        self.successes.append(message)
        
    def print_error(self, message):
        print(f"{Fore.RED}✗ {message}{Style.RESET_ALL}")
        self.errors.append(message)
        
    def print_warning(self, message):
        print(f"{Fore.YELLOW}⚠ {message}{Style.RESET_ALL}")
        self.warnings.append(message)
        
    def print_info(self, message):
        print(f"{Fore.BLUE}ℹ {message}{Style.RESET_ALL}")
        
    def print_data(self, data, indent=2):
        """Print data in pretty format"""
        if isinstance(data, (dict, list)):
            print(" " * indent + json.dumps(data, indent=2, ensure_ascii=False))
        else:
            print(" " * indent + str(data))

    def test_backend_connection(self):
        """Test backend server connection"""
        self.print_header("Testing Backend Connection")
        
        try:
            response = requests.get(f"{BASE_URL}/docs", timeout=5)
            if response.status_code == 200:
                self.print_success(f"Server is running at {BASE_URL}")
                return True
            else:
                self.print_error(f"Server responding with status code {response.status_code}")
                return False
        except requests.exceptions.ConnectionError:
            self.print_error(f"Cannot connect to server at {BASE_URL}")
            print(f"  {Fore.YELLOW}Make sure server is running: uvicorn app.main:app --reload")
            return False
        except Exception as e:
            self.print_error(f"Connection error: {e}")
            return False

    def login(self):
        """Fake login - use dummy token"""
        self.print_header("System Authentication (Dummy Mode)")
        try:
            self.token = "dummy_token_for_tests_123456"
            self.print_success(f"Dummy login successful - Token: {self.token}")
            return True
        except Exception as e:
            self.print_error(f"Dummy login failed: {e}")
            return False

    def register_and_login(self):
        """Skip real registration - not needed"""
        self.print_info("Skipping user registration (dummy mode)")
        return self.login()

    def get_headers(self):
        """Get API request headers"""
        return {
            "Authorization": f"Bearer {self.token}",
            "Content-Type": "application/json"
        }

    def test_recommendations_endpoints(self):
        """Test all recommendations endpoints"""
        self.print_header("Testing Recommendations Endpoints")
        
        # 1. Test GET all recommendations
        self.print_info("Testing GET /api/v1/recommendations/")
        try:
            response = requests.get(
                f"{API_URL}/recommendations/",
                headers=self.get_headers()
            )
            
            if response.status_code == 200:
                recommendations = response.json()
                self.print_success(f"Received {len(recommendations)} recommendations")
                
                if len(recommendations) > 0:
                    self.print_info("First recommendation sample:")
                    self.print_data(recommendations[0])
                else:
                    self.print_warning("No recommendations in system")
                    
                return recommendations
            else:
                self.print_error(f"Failed to get recommendations: {response.status_code}")
                self.print_data(response.json())
                return []
                
        except Exception as e:
            self.print_error(f"Error requesting recommendations: {e}")
            return []

    def test_upload_recommendation(self):
        """Test recommendation file upload"""
        self.print_header("Testing File Upload")
        
        # Create test Word file
        test_file_path = "test_recommendations.docx"
        
        try:
            # Check if python-docx is available
            try:
                from docx import Document
                
                doc = Document()
                doc.add_heading('Nutritional Recommendations', 0)
                doc.add_heading('Home Recommendations', level=1)
                doc.add_paragraph('1. Eat more vegetables')
                doc.add_paragraph('2. Drink at least 8 glasses of water per day')
                doc.add_paragraph('3. Reduce processed food')
                doc.save(test_file_path)
                
                self.print_success(f"Test file created: {test_file_path}")
                
            except ImportError:
                self.print_warning("python-docx not installed, using dummy file")
                # Create simple text file
                with open(test_file_path, 'w') as f:
                    f.write("Home Recommendations\n1. First recommendation\n2. Second recommendation")
                test_file_path = "test_recommendations.txt"
                
            # Upload file
            with open(test_file_path, 'rb') as f:
                files = {'file': ('recommendations.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
                data = {'visit_date': datetime.now().strftime('%Y-%m-%d')}
                
                response = requests.post(
                    f"{API_URL}/recommendations/upload",
                    headers={"Authorization": f"Bearer {self.token}"},
                    files=files,
                    data=data
                )
                
                if response.status_code == 200:
                    result = response.json()
                    self.print_success("File uploaded successfully!")
                    self.print_data(result)
                    return result
                else:
                    self.print_error(f"Upload failed: {response.status_code}")
                    self.print_data(response.text)
                    return None
                    
        except Exception as e:
            self.print_error(f"File upload error: {e}")
            return None
        finally:
            # Cleanup
            if os.path.exists(test_file_path):
                os.remove(test_file_path)

    def test_compliance_endpoints(self):
        """Test compliance endpoints"""
        self.print_header("Testing Compliance Endpoints")
        
        # First get recommendations
        recommendations = self.test_recommendations_endpoints()
        
        if not recommendations:
            self.print_warning("No recommendations, cannot test compliance")
            return
            
        # Test getting compliance report
        rec_id = recommendations[0]['id']
        visit_date = recommendations[0]['visit_date']
        
        # Calculate two week period
        start_date = datetime.strptime(visit_date, '%Y-%m-%d')
        end_date = start_date + timedelta(days=14)
        period = f"{visit_date} to {end_date.strftime('%Y-%m-%d')}"
        
        self.print_info(f"Testing compliance for period: {period}")
        
        try:
            response = requests.get(
                f"{API_URL}/compliance/report",
                headers=self.get_headers(),
                params={
                    'recommendation_id': rec_id,
                    'period': period
                }
            )
            
            if response.status_code == 200:
                report = response.json()
                self.print_success("Compliance report received!")
                self.print_data(report)
                return report
            else:
                self.print_error(f"Failed to get report: {response.status_code}")
                self.print_data(response.text)
                return None
                
        except Exception as e:
            self.print_error(f"Compliance request error: {e}")
            return None

    def check_frontend_files(self):
        """Check frontend files"""
        self.print_header("Checking Frontend Files")
        
        # Critical files to check
        critical_files = [
            "frontend/src/pages/RecommendationsPage.js",
            "frontend/src/pages/CompliancePage.js",
            "frontend/src/services/recommendationService.js",
            "frontend/src/services/complianceService.js"
        ]
        
        missing_files = []
        
        for file in critical_files:
            if os.path.exists(file):
                self.print_success(f"File exists: {file}")
                
                # Check file size
                size = os.path.getsize(file)
                if size == 0:
                    self.print_warning(f"  File is empty!")
                else:
                    self.print_info(f"  Size: {size} bytes")
            else:
                self.print_error(f"File missing: {file}")
                missing_files.append(file)
                
        return missing_files

    def check_database_tables(self):
        """Check database tables"""
        self.print_header("Checking Database")
        
        if not os.path.exists("backend/nutrition_tracker.db"):
            self.print_error("Database file not found: nutrition_tracker.db")
            return
            
        try:
            import sqlite3
            conn = sqlite3.connect("backend/nutrition_tracker.db")
            cursor = conn.cursor()
            
            # Check tables
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            
            self.print_info(f"Found {len(tables)} tables:")
            for table in tables:
                self.print_success(f"  - {table[0]}")
                
                # Count records
                cursor.execute(f"SELECT COUNT(*) FROM {table[0]}")
                count = cursor.fetchone()[0]
                self.print_info(f"    Records: {count}")
                
            # Check recommendations table specifically
            cursor.execute("SELECT * FROM nutritionist_recommendations LIMIT 1")
            if cursor.fetchone():
                self.print_success("Data exists in recommendations table")
            else:
                self.print_warning("No data in recommendations table")
                
            conn.close()
            
        except Exception as e:
            self.print_error(f"Database access error: {e}")

    def test_cors_configuration(self):
        """Test CORS configuration"""
        self.print_header("Testing CORS Configuration")
        
        try:
            # Test OPTIONS request
            response = requests.options(
                f"{API_URL}/recommendations/",
                headers={
                    "Origin": FRONTEND_URL,
                    "Access-Control-Request-Method": "GET"
                }
            )
            
            cors_headers = response.headers.get("Access-Control-Allow-Origin")
            if cors_headers:
                self.print_success(f"CORS configured: {cors_headers}")
            else:
                self.print_warning("CORS headers not found")
                
        except Exception as e:
            self.print_error(f"CORS test error: {e}")

    def run_full_diagnostic(self):
        """Run all tests"""
        print(f"\n{Fore.MAGENTA}╔{'═'*58}╗")
        print(f"{Fore.MAGENTA}║{' '*15}RECOMMENDATIONS SYSTEM DIAGNOSTIC{' '*10}║")
        print(f"{Fore.MAGENTA}╚{'═'*58}╝{Style.RESET_ALL}\n")
        
        # 1. Test server
        if not self.test_backend_connection():
            self.print_error("\nServer not running - cannot continue tests")
            return
            
        # 2. Login (dummy)
        if not self.login():
            self.print_error("\nAuthentication failed - cannot continue")
            return
            
        # 3. Test endpoints
        self.test_recommendations_endpoints()
        
        # 4. Test file upload
        self.test_upload_recommendation()
        
        # 5. Test compliance
        self.test_compliance_endpoints()
        
        # 6. Test CORS
        self.test_cors_configuration()
        
        # 7. Check files
        self.check_frontend_files()
        
        # 8. Check database
        self.check_database_tables()
        
        # Summary
        self.print_summary()

    def print_summary(self):
        """Print test summary"""
        self.print_header("TEST SUMMARY")
        
        total = len(self.successes) + len(self.errors) + len(self.warnings)
        
        print(f"{Fore.GREEN}Successes: {len(self.successes)}/{total}")
        print(f"{Fore.YELLOW}Warnings: {len(self.warnings)}/{total}")
        print(f"{Fore.RED}Errors: {len(self.errors)}/{total}")
        
        if self.errors:
            print(f"\n{Fore.RED}Critical errors found:")
            for error in self.errors:
                print(f"  • {error}")
                
        if self.warnings:
            print(f"\n{Fore.YELLOW}Warnings:")
            for warning in self.warnings:
                print(f"  • {warning}")
                
        # Recommendations for fixes
        print(f"\n{Fore.CYAN}Recommendations for fixes:")
        
        if "recommendationService.js" in str(self.errors):
            print("  1. Create the file src/services/recommendationService.js")
            
        if "CORS" in str(self.warnings):
            print("  2. Check CORS settings in main.py")
            
        if "No data in recommendations table" in str(self.warnings):
            print("  3. Upload recommendations file through the interface")
            
        print(f"\n{Fore.GREEN}Test complete!{Style.RESET_ALL}")

def main():
    """Main function"""
    tester = TestRunner()
    
    try:
        tester.run_full_diagnostic()
    except KeyboardInterrupt:
        print(f"\n{Fore.YELLOW}Test interrupted by user{Style.RESET_ALL}")
    except Exception as e:
        print(f"\n{Fore.RED}General error: {e}{Style.RESET_ALL}")

if __name__ == "__main__":
    main()
